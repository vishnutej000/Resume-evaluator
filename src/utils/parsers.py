import fitz
import docx
import re
import io
import os
import pandas as pd

class ResumeParser:
    def __init__(self):
        self.email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        self.github_pattern = r'github\.com/([a-zA-Z0-9_-]+)'
        self.max_file_size = 10 * 1024 * 1024  # 10MB
        self.supported_formats = ['.pdf', '.docx']

    def validate_file(self, content, filename):
        if len(content.getvalue()) > self.max_file_size:
            print(f"[DEBUG] File {filename} exceeds 10MB limit")
            return f"File {filename} exceeds 10MB limit"
        file_format = os.path.splitext(filename)[1].lower()
        if file_format not in self.supported_formats:
            print(f"[DEBUG] Unsupported file format: {file_format}")
            return f"Unsupported file format: {file_format}"
        try:
            if file_format == '.pdf':
                fitz.open(stream=content.getvalue(), filetype="pdf")
            else:
                docx.Document(content)
        except Exception as e:
            print(f"[DEBUG] File {filename} appears to be corrupted: {str(e)}")
            return f"File {filename} appears to be corrupted: {str(e)}"
        return None

    def extract_from_pdf(self, content):
        try:
            doc = fitz.open(stream=content.getvalue(), filetype="pdf")
            if doc.is_encrypted:
                print("[DEBUG] PDF is password-protected")
                return "", "Password-protected PDF"
            text = ""
            for page in doc:
                page_text = page.get_text("text", sort=True)
                if not page_text:
                    print("[DEBUG] No text found on a PDF page")
                text += page_text
                tables = page.find_tables()
                for table in tables:
                    df = table.to_pandas()
                    text += "\n" + df.to_string()
            doc.close()
            if not text.strip():
                print("[DEBUG] No text extracted from PDF")
                return "", "No text extracted from PDF"
            return text, None
        except Exception as e:
            print(f"[DEBUG] PDF extraction failed: {e}")
            return "", f"PDF extraction failed: {e}"

    def extract_from_docx(self, content):
        try:
            doc = docx.Document(content)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text for cell in row.cells)
                    text += row_text + "\n"
            if not text.strip():
                print("[DEBUG] No text extracted from DOCX")
                return "", "No text extracted from DOCX"
            return text, None
        except Exception as e:
            print(f"[DEBUG] DOCX extraction failed: {e}")
            return "", f"DOCX extraction failed: {e}"

    def clean_text(self, text):
        text = re.sub(r'[^\w\s@.-]', ' ', text)
        text = re.sub(r'\s+', ' ', text)
        return text.strip()

    def extract_email(self, text):
        emails = re.findall(self.email_pattern, text)
        return emails[0] if emails else ""

    def extract_github(self, text):
        matches = re.findall(self.github_pattern, text.lower())
        return matches[0] if matches else ""

    def extract_skills(self, text):
        common_skills = [
            'Python', 'Java', 'JavaScript', 'TypeScript', 'C++', 'C#', 'Ruby', 'PHP', 'Go', 'Rust',
            'React', 'Angular', 'Vue.js', 'Node.js', 'Express', 'Django', 'Flask', 'Spring', 'ASP.NET',
            'SQL', 'MySQL', 'PostgreSQL', 'MongoDB', 'Redis', 'Cassandra', 'Oracle',
            'AWS', 'Azure', 'GCP', 'Docker', 'Kubernetes', 'Terraform', 'Ansible', 'Jenkins', 'GitLab CI',
            'Git', 'SVN', 'JIRA', 'Confluence',
            'Machine Learning', 'Deep Learning', 'TensorFlow', 'PyTorch', 'Scikit-learn', 'NLP', 'Computer Vision',
            'REST API', 'GraphQL', 'gRPC', 'WebSocket', 'Microservices', 'Serverless', 'CI/CD'
        ]
        found_skills = []
        text_lower = text.lower()
        for skill in common_skills:
            if skill.lower() in text_lower:
                found_skills.append(skill)
        return found_skills

    def extract_experience(self, text):
        patterns = [
            r'(\d+)\+?\s*years?\s*(?:of\s*)?experience',
            r'experience.*?(\d+)\+?\s*years?',
            r'(\d+)\+?\s*yrs?\s*experience',
            r'(\d+)\+?\s*years?\s*in\s*the\s*field',
            r'(\d+)\+?\s*years?\s*of\s*professional',
            r'(\d+)\+?\s*years?\s*working\s*experience'
        ]
        for pattern in patterns:
            matches = re.findall(pattern, text.lower())
            if matches:
                try:
                    return int(matches[0])
                except ValueError:
                    continue
        return 0

    def extract_education(self, text):
        degrees = [
            'B.Tech', 'M.Tech', 'B.S.', 'M.S.', 'PhD', 'Bachelor', 'Master',
            'B.E.', 'M.E.', 'B.Sc.', 'M.Sc.', 'B.A.', 'M.A.',
            'B.Com', 'M.Com', 'MBA', 'BBA', 'MCA', 'BCA'
        ]
        text_lower = text.lower()
        for degree in degrees:
            if degree.lower() in text_lower:
                return degree
        return ""

    def parse_resume(self, filename, content):
        print(f"[DEBUG] Parsing {filename}")
        error = self.validate_file(content, filename)
        if error:
            print(f"[DEBUG] Validation error: {error}")
            return {'filename': filename, 'error': error}
        if filename.lower().endswith('.pdf'):
            text, error = self.extract_from_pdf(content)
        elif filename.lower().endswith('.docx'):
            text, error = self.extract_from_docx(content)
        else:
            print(f"[DEBUG] Unsupported file type for {filename}")
            return {'filename': filename, 'error': 'Unsupported file type'}
        if error:
            print(f"[DEBUG] Extraction error: {error}")
            return {'filename': filename, 'error': error}
        if not text:
            print(f"[DEBUG] No text extracted from {filename}")
            return {'filename': filename, 'error': 'No text extracted'}
        text = self.clean_text(text)
        print(f"[DEBUG] Cleaned text length for {filename}: {len(text)}")
        return {
            'filename': filename,
            'text': text,
            'email': self.extract_email(text),
            'github_username': self.extract_github(text),
            'skills': self.extract_skills(text),
            'experience_years': self.extract_experience(text),
            'education': self.extract_education(text),
            'error': None
        }